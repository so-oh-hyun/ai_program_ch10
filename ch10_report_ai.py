import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# MS워드 문서 변환 함수 정의
def markdown_to_docx(markdown_content: str, font_name: str, base_font_size: int):
	doc = Document()
	lines = markdown_content.split("\n")
	for line in lines:
		line = line.strip()
		if not line:
			continue
		if line.startswith("## "):
			paragraph = doc.add_paragraph()
			run = paragraph.add_run(line[3:])
			font = run.font
			font.size = Pt(base_font_size + 3)
			font.name = font_name
			font.bold = True
			run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
		elif line.startswith("### "):
			paragraph = doc.add_paragraph()
			run = paragraph.add_run(line[4:])
			font = run.font
			font.size = Pt(base_font_size + 1)
			font.name = font_name
			font.bold = True
			run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
		else:
			paragraph = doc.add_paragraph()
			run = paragraph.add_run(line)
			font = run.font
			font.size = Pt(base_font_size)
			font.name = font_name
			run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
	byte_io = BytesIO()
	doc.save(byte_io)
	byte_io.seek(0)
	return byte_io

def main():
	st.set_page_config(layout="wide")
	st.title("보고서 작성 프로그램")
	with st.sidebar:
		openai_api_key = st.text_input("OpenAI API Key", type="password")
		# 폰트 종류 및 크기 선택 위젯 추가
		font_name = st.selectbox("글꼴 선택:", ["맑은 고딕", "바탕체"])
		base_font_size = st.slider("기본 글자 크기 (pt):", value=11)
		# OpenAI 클라이언트 생성
		if openai_api_key:
			client = OpenAI(api_key=openai_api_key)
	# 보고서 작성 함수 정의
	def process_text(prompt, text):
		content = prompt + "\n" + text
		response = client.chat.completions.create(
			model="gpt-4o-mini",
			messages=[{"role": "user", "content": content}],
		)
		return response.choices[0].message.content
	# 보고서 작성을 위한 프롬프트 입력
	prompt = """
	너는 보고서 작성 전문가야.
	다음 형식으로 보고서를 작성해줘.
	- 마크다운을 활용해 체계적으로 작성할 것
	- heading2(##) 3개, 각 heading2 내에서는 heading3(###) 2개로 구성할 것
	- heading2의 내용은 300자 이상으로 작성할 것
	- 목차는 제외할 것
	- 보고서 내용만 응답 결과로 보여줄 것
	"""
	default_user_input = """생성형 AI가 세상을 어떻게 바꿀 수 있을까?"""
	user_input = st.text_area(
		"작성할 보고서의 주제 또는 내용을 입력하세요:",
		value=default_user_input,
		height=70,
	)
	# 보고서 작성
	if st.button("보고서 작성"):
		if not openai_api_key:
			st.info("계속하려면 OpenAI API Key를 추가하세요.")
			st.stop()
		if not user_input.strip():
			st.warning("작성할 보고서의 주제를 입력하세요.")
			st.stop()
		with st.spinner("작성 중..."):
			result = process_text(prompt, user_input)
			st.write(result)
			# MS워드 문서 변환 함수 호출
			docx_file = markdown_to_docx(result, font_name, base_font_size)
			# 다운로드 버튼 생성
			st.download_button(
				label="보고서 다운로드",
				data=docx_file,
				file_name="보고서.docx",
			)

if __name__ == "__main__":
	main()
